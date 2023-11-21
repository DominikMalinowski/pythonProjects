
# creating document object 
import docx
doc = docx.Document('demo.docx')

# getting text from paragraph
import docx
doc = docx.Document('demo.docx')
len(doc.paragraphs)
doc.paragraph[0].text

# getting text from runs
import docx
doc = docx.Document('demo.docx')
len(doc.paragraphs[1].runs)
doc.paragraph[1].runs[3].text

# adding paragraph and runs to document
import docx
doc = docx.Document('demo.docx')
paragraph_1 = doc.add_paragraph('First paragraph')
paragraph_2 = doc.add_paragraph('Second paragraph')

paragraph_1.add_run('First run')
paragraph_2.add_run('Second run')

# adding headings (with it levels 0-4)
import docx
doc = docx.Document('demo.docx')
doc.add_heading('Title','0') 
doc.add_heading('Header_1','1') 
doc.add_heading('Subheader_2','2') 

# text attributes and style 
import docx
doc = docx.Document('demo.docx')
doc.paragraph[0].runs[1].style ='Heading1' 
doc.paragraph[0].runs[1].bold=True 
doc.paragraph[0].runs[1].italic=True 
doc.paragraph[0].runs[1].underline=True 
doc.paragraph[0].runs[1].strike=True 
doc.paragraph[0].runs[1].double_strike=True 
doc.paragraph[0].runs[1].all_caps=True 
doc.paragraph[0].runs[1].small_caps=True 
doc.paragraph[0].runs[1].shadow=True 
doc.paragraph[0].runs[1].outline=True 
doc.paragraph[0].runs[1].rtl=True 
doc.paragraph[0].runs[1].imprint=True 
doc.paragraph[0].runs[1].emboss=True 

# adding page break  
import docx
doc = docx.Document('demo.docx')
doc.add_break(doc.text.WD_BREAK_PAGE)

# adding image 
import docx
doc = docx.Document('demo.docx')
doc.add_picture('picture.png', width=docx.shared.Inches(1), height=docx.shared.CM(4))

# saving document
import docx
doc = docx.Document('demo.docx')
doc.add_paragraph('Hello')
doc.save('changed_file.docx')

# getting whole document 
import docx

def get_text(filename):
    doc = docx.Document(filename)
    fullText = []
    for paragraph in doc.paragraph:
        fullText.append(paragraph.text)
    return '\n'.join(fullText)